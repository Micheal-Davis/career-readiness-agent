from docx import Document
import pytest

from career_agent.initialization import (
    initialize_empty_profile,
    initialize_profile_from_resume,
)
from career_agent.profile import (
    CapabilityProfileStore,
    ConfirmationStatus,
    EvidenceType,
)


@pytest.fixture
def store(tmp_path):
    database = CapabilityProfileStore(tmp_path / "profiles.sqlite3")
    yield database
    database.close()


def test_empty_and_resume_paths_each_create_one_profile(store, tmp_path):
    empty = initialize_empty_profile(store)
    resume = tmp_path / "resume.txt"
    resume.write_text("项目经历\n职业 Agent\n用 Python 实现检索服务", encoding="utf-8")

    imported = initialize_profile_from_resume(
        store,
        resume_path=resume,
        source_documents_dir=tmp_path / "source-documents",
    )

    assert empty.profile.id != imported.profile.id
    assert imported.drafts[0].confirmation_status == ConfirmationStatus.DRAFT
    assert imported.drafts[0].evidence_type == EvidenceType.PROJECT
    assert imported.drafts[0].details["extracted_text"] == "职业 Agent\n用 Python 实现检索服务"


def test_uploaded_docx_is_copied_and_remains_a_source_document(store, tmp_path):
    resume = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("竞赛经历")
    document.add_paragraph("数学建模竞赛，省二等奖")
    document.save(resume)

    result = initialize_profile_from_resume(
        store,
        resume_path=resume,
        source_documents_dir=tmp_path / "source-documents",
    )

    copied_path = tmp_path / "source-documents" / result.profile.id / "resume.docx"
    assert copied_path.read_bytes() == resume.read_bytes()
    assert result.drafts[0].source_document_id is not None
    assert result.drafts[0].evidence_type == EvidenceType.COMPETITION


def test_user_can_edit_confirm_or_discard_extracted_drafts(store, tmp_path):
    resume = tmp_path / "resume.txt"
    resume.write_text("项目经历\n职业 Agent", encoding="utf-8")
    result = initialize_profile_from_resume(
        store,
        resume_path=resume,
        source_documents_dir=tmp_path / "source-documents",
    )
    draft = result.drafts[0]

    edited = store.update_draft_evidence_record(
        draft.id,
        title="职业 Agent",
        details={
            "technologies": "Python, FastAPI",
            "contribution": "设计 API 与检索流程",
            "result": "完成可演示原型",
        },
    )
    confirmed = store.confirm_evidence_record(edited.id)

    assert confirmed.confirmation_status == ConfirmationStatus.CONFIRMED
    with pytest.raises(ValueError, match="Only draft"):
        store.discard_draft_evidence_record(confirmed.id)

    another = store.add_evidence_record(
        result.profile.id,
        evidence_type=EvidenceType.CAMPUS,
        title="校园经历",
        details={"extracted_text": "学生组织负责人"},
    )
    store.discard_draft_evidence_record(another.id)
    assert store.list_draft_evidence(result.profile.id) == []


def test_unfinished_draft_cannot_be_confirmed(store, tmp_path):
    resume = tmp_path / "resume.txt"
    resume.write_text("项目经历\n职业 Agent", encoding="utf-8")
    result = initialize_profile_from_resume(
        store,
        resume_path=resume,
        source_documents_dir=tmp_path / "source-documents",
    )

    with pytest.raises(ValueError, match="technologies"):
        store.confirm_evidence_record(result.drafts[0].id)


def test_resume_parser_splits_adjacent_dated_projects_and_cleans_competition_title(store, tmp_path):
    resume = tmp_path / "resume.txt"
    resume.write_text(
        """项目经历
《基于多标注的可解释动作评估》 第一作者 2025.09-2026.04
项目内容：构建精准且可解释的评分模型。
1. 提出检测注意力机制聚焦运动员，并按阶段进行因果推理。
2. 在 FineDiving 数据集上 SRCC 达 0.9545，显著提升评分一致性。
计算机网络核心协议栈与分布式Web架构实战 开发成员 2026.03-2026.06
技术栈：Python, Raw Socket, 多线程/异步编程, HTTP 协议
项目描述：从零构建网络诊断工具与分布式 Web 服务系统。
核心工作与成果：实现代理缓存，降低响应延迟。

竞赛经历
全国大学生数学建模竞赛 队长 省级二等奖
个人贡献：负责建模与论文撰写
""",
        encoding="utf-8",
    )

    result = initialize_profile_from_resume(
        store, resume_path=resume, source_documents_dir=tmp_path / "sources"
    )

    assert [(item.evidence_type, item.title) for item in result.drafts] == [
        (EvidenceType.PROJECT, "基于多标注的可解释动作评估"),
        (EvidenceType.PROJECT, "计算机网络核心协议栈与分布式Web架构实战"),
        (EvidenceType.COMPETITION, "全国大学生数学建模竞赛"),
    ]
    assert result.drafts[0].details["result"] == "在 FineDiving 数据集上 SRCC 达 0.9545，显著提升评分一致性。"
    assert result.drafts[1].details["technologies"] == "Python, Raw Socket, 多线程/异步编程, HTTP 协议"
    assert result.drafts[2].details["outcome"] == "省级二等奖"
    assert "organizer" not in result.drafts[2].details
