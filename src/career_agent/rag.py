"""Document loading, chunking, and Chroma vector indexing."""
from __future__ import annotations

from pathlib import Path
import shutil
from uuid import uuid4

from docx import Document as WordDocument
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from .embeddings import HashEmbeddings


def _read_docx(path: Path) -> str:
    """Extract non-empty paragraphs and table rows from a .docx document."""
    word_document = WordDocument(path)

    paragraphs = [
        paragraph.text.strip()
        for paragraph in word_document.paragraphs
        if paragraph.text.strip()
    ]

    table_rows = [
        " | ".join(cell.text.strip() for cell in row.cells)
        for table in word_document.tables
        for row in table.rows
    ]

    return "\n".join(paragraphs + table_rows)


def load_documents(documents_dir: Path) -> list[Document]:
    """Load .md, .txt, and .docx files and retain each source file name."""
    documents: list[Document] = []

    for path in sorted(documents_dir.rglob("*")):
        if not path.is_file():
            continue

        suffix = path.suffix.lower()

        if suffix in {".md", ".txt"}:
            content = path.read_text(encoding="utf-8")
        elif suffix == ".docx":
            content = _read_docx(path)
        else:
            continue

        if content.strip():
            documents.append(
                Document(
                    page_content=content,
                    metadata={
                        "source": path.name,
                        "source_path": path.relative_to(documents_dir).as_posix(),
                    },
                )
            )

    if not documents:
        raise ValueError(
            f"No .md, .txt, or .docx documents found in {documents_dir}"
        )

    return documents


def build_index(documents_dir: Path, persist_dir: Path) -> int:
    """Rebuild the Chroma index from the current source documents only.

    Source documents are loaded before removing an existing index so a failed
    load does not discard the last usable index.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=80,
    )
    chunks = splitter.split_documents(load_documents(documents_dir))

    chunk_counts: dict[str, int] = {}
    for chunk in chunks:
        source_path = chunk.metadata["source_path"]
        chunk_counts[source_path] = chunk_counts.get(source_path, 0) + 1
        chunk.metadata["chunk_index"] = chunk_counts[source_path]

    version_dir = persist_dir.with_name(f"{persist_dir.name}.version-{uuid4().hex}")
    pointer_path = persist_dir.with_name(f"{persist_dir.name}.active")
    temporary_pointer_path = pointer_path.with_name(f"{pointer_path.name}.tmp-{uuid4().hex}")
    try:
        Chroma.from_documents(
            documents=chunks,
            embedding=HashEmbeddings(),
            persist_directory=str(version_dir),
            collection_name="career_knowledge",
        )
        temporary_pointer_path.write_text(version_dir.name, encoding="utf-8")
        temporary_pointer_path.replace(pointer_path)
    except Exception:
        temporary_pointer_path.unlink(missing_ok=True)
        raise

    return len(chunks)


def get_retriever(persist_dir: Path):
    """Open the existing Chroma index as a retriever."""
    store = Chroma(
        collection_name="career_knowledge",
        persist_directory=str(active_index_directory(persist_dir)),
        embedding_function=HashEmbeddings(),
    )

    return store.as_retriever(search_kwargs={"k": 4})


def active_index_directory(persist_dir: Path) -> Path:
    """Return the current versioned index, retaining compatibility with legacy indexes."""
    pointer_path = persist_dir.with_name(f"{persist_dir.name}.active")
    if pointer_path.is_file():
        active_name = pointer_path.read_text(encoding="utf-8").strip()
        if not active_name:
            raise ValueError("No active knowledge-base index is available.")
        return pointer_path.with_name(active_name)
    return persist_dir


def clear_active_index(persist_dir: Path) -> None:
    """Mark an empty document library as having no searchable index."""
    pointer_path = persist_dir.with_name(f"{persist_dir.name}.active")
    pointer_path.write_text("", encoding="utf-8")
