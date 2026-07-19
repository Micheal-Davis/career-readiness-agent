"""First-run capability-profile initialization from an empty form or resume."""
from __future__ import annotations

import shutil
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document as WordDocument
from pypdf import PdfReader

from .profile import CapabilityProfile, CapabilityProfileStore, EvidenceRecord, EvidenceType
from .analysis import ResumeExtraction, ResumeExtractor

_SUPPORTED_SUFFIXES = {".docx", ".pdf", ".txt", ".md"}
_SECTION_TYPES: tuple[tuple[EvidenceType, tuple[str, ...]], ...] = (
    (EvidenceType.PROJECT, ("项目经历", "项目经验", "项目")),
    (EvidenceType.WORK, ("实习经历", "工作经历", "实习", "工作经验")),
    (EvidenceType.COMPETITION, ("竞赛经历", "竞赛", "获奖经历", "获奖")),
    (EvidenceType.COURSE, ("课程", "学习经历", "教育经历")),
    (EvidenceType.CAMPUS, ("校园经历", "社团经历", "学生工作")),
)
_DATE_RANGE_PATTERN = re.compile(r"\d{4}[./]\d{1,2}\s*[-—~至]\s*\d{4}[./]\d{1,2}")
_ROLE_PATTERN = re.compile(r"\s*(第一作者|开发成员|队长|队员|成员|负责人|组长|核心成员|项目负责人)\s*")
_AWARD_PATTERN = re.compile(r"(国家级|省级|校级|区域赛|全国赛)?[一二三四五六七八九十]等奖|金奖|银奖|铜奖|优秀奖")


@dataclass(frozen=True)
class ProfileInitialization:
    """One profile and its unconfirmed evidence candidates."""

    profile: CapabilityProfile
    drafts: tuple[EvidenceRecord, ...]


def initialize_empty_profile(store: CapabilityProfileStore) -> ProfileInitialization:
    """Start the same confirmation flow without an uploaded resume."""
    return ProfileInitialization(profile=store.create_profile(), drafts=())


def initialize_profile_from_resume(
    store: CapabilityProfileStore,
    *,
    resume_path: Path,
    source_documents_dir: Path,
    resume_extractor: ResumeExtractor | None = None,
) -> ProfileInitialization:
    """Copy an uploaded resume, then make editable evidence drafts from its text."""
    if not resume_path.is_file():
        raise ValueError("Uploaded resume does not exist.")
    if resume_path.suffix.lower() not in _SUPPORTED_SUFFIXES:
        supported = ", ".join(sorted(_SUPPORTED_SUFFIXES))
        raise ValueError(f"Unsupported resume format. Use one of: {supported}.")

    resume_text = _read_resume(resume_path)
    extracted = (
        _model_extracted_evidence(resume_extractor(resume_text))
        if resume_extractor is not None
        else _extract_evidence_sections(resume_text)
    )

    profile = store.create_profile()
    destination = source_documents_dir / profile.id / resume_path.name
    destination.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(resume_path, destination)
    document = store.add_source_document(
        profile.id,
        filename=resume_path.name,
        source_path=destination.as_posix(),
    )

    drafts = [
        store.add_evidence_record(
            profile.id,
            evidence_type=evidence_type,
            title=title,
            details=details,
            source_document_id=document.id,
        )
        for evidence_type, title, details in extracted
    ]
    return ProfileInitialization(profile=profile, drafts=tuple(drafts))


def _model_extracted_evidence(
    result: list[dict] | ResumeExtraction,
) -> list[tuple[EvidenceType, str, dict[str, str]]]:
    extraction = ResumeExtraction.model_validate({"evidence": result} if isinstance(result, list) else result)
    return [
        (item.evidence_type, item.title, {**item.details, "extracted_text": "模型结构化提取"})
        for item in extraction.evidence
        if item.title.strip()
    ]


def _read_resume(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8")
    if suffix == ".docx":
        document = WordDocument(path)
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
        table_rows = [
            " | ".join(cell.text.strip() for cell in row.cells)
            for table in document.tables
            for row in table.rows
        ]
        return "\n".join(line for line in paragraphs + table_rows if line)
    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_evidence_sections(text: str) -> list[tuple[EvidenceType, str, dict[str, str]]]:
    """Split each conventional resume section into separately editable experience drafts."""
    sections: list[tuple[EvidenceType, str, list[str]]] = []
    current: tuple[EvidenceType, str, list[str]] | None = None

    for raw_line in text.splitlines():
        line = raw_line.strip()
        evidence_type = _section_type_for(line)
        if evidence_type is not None:
            if current is not None:
                sections.append(current)
            current = (evidence_type, line, [])
        elif current is not None:
            current[2].append(line)

    if current is not None:
        sections.append(current)

    return [
        (evidence_type, title, _prefill_details(evidence_type, item_lines))
        for evidence_type, _section_title, lines in sections
        for title, item_lines in _split_section_items(evidence_type, lines)
    ]


def _split_section_items(
    evidence_type: EvidenceType, lines: list[str]
) -> list[tuple[str, list[str]]]:
    """Split experiences by blank lines or dated title rows within a single resume section."""
    blocks: list[list[str]] = []
    current: list[str] = []
    for line in lines:
        is_dated_item_start = (
            bool(current)
            and _DATE_RANGE_PATTERN.search(line) is not None
            and "：" not in line
            and ":" not in line
        )
        if not line or is_dated_item_start:
            if current:
                blocks.append(current)
                current = []
            if not line:
                continue
        current.append(line)
    if current:
        blocks.append(current)
    return [(_clean_item_title(evidence_type, block[0]), block) for block in blocks if block]


def _clean_item_title(evidence_type: EvidenceType, value: str) -> str:
    title = value.lstrip("-•· ").split("|", maxsplit=1)[0].strip()
    quoted = re.search(r"《([^》]+)》", title)
    if quoted:
        return quoted.group(1).strip()
    if evidence_type == EvidenceType.COMPETITION:
        match = re.search(r".+?竞赛", title)
        if match:
            return match.group(0).strip()
    title = _DATE_RANGE_PATTERN.sub("", title)
    title = _ROLE_PATTERN.sub(" ", title)
    title = _AWARD_PATTERN.sub("", title)
    return " ".join(title.split())


def _prefill_details(evidence_type: EvidenceType, lines: list[str]) -> dict[str, str]:
    """Extract labelled resume fields without inventing missing achievements or skills."""
    values: dict[str, str] = {"extracted_text": "\n".join(lines)}
    labels = {
        EvidenceType.PROJECT: {
            "technologies": ("技术", "技术栈", "工具"),
            "contribution": ("职责", "个人贡献", "负责", "工作内容"),
            "result": ("成果", "结果", "项目成果", "效果"),
        },
        EvidenceType.WORK: {
            "responsibilities": ("职责", "工作内容", "负责"),
            "work_content": ("工作内容", "成果", "结果", "效果"),
        },
        EvidenceType.COMPETITION: {
            "outcome": ("奖项", "获奖", "成绩", "结果"),
            "contribution": ("个人贡献", "职责", "负责"),
        },
        EvidenceType.COURSE: {
            "course_or_activity": ("课程", "活动"),
            "outcome": ("成果", "结果", "成绩"),
            "related_work": ("实践", "作业", "项目"),
        },
        EvidenceType.CAMPUS: {
            "responsibilities": ("职责", "负责", "工作内容"),
            "developed_capabilities": ("锻炼能力", "能力", "收获", "成果", "结果"),
        },
    }[evidence_type]
    contributions: list[str] = []
    results: list[str] = []
    for line in lines[1:]:
        if "：" not in line and ":" not in line:
            if evidence_type == EvidenceType.PROJECT and re.match(r"\d+[.、]", line):
                cleaned_line = re.sub(r"^\d+[.、]\s*", "", line)
                (results if _looks_like_result(line) else contributions).append(cleaned_line)
            continue
        label, value = line.replace(":", "：", 1).split("：", maxsplit=1)
        for field, aliases in labels.items():
            if label.strip() in aliases and value.strip():
                values[field] = value.strip()
                break
        else:
            if evidence_type == EvidenceType.PROJECT and label.strip() in {"项目内容", "项目描述", "核心工作"}:
                contributions.append(value.strip())
            elif evidence_type == EvidenceType.PROJECT and label.strip() in {"核心工作与成果"}:
                results.append(value.strip())
    if evidence_type == EvidenceType.PROJECT:
        if contributions and not values.get("contribution"):
            values["contribution"] = "\n".join(contributions)
        if results and not values.get("result"):
            values["result"] = "\n".join(results)
    if evidence_type == EvidenceType.COMPETITION:
        award = _AWARD_PATTERN.search(lines[0])
        if award and not values.get("outcome"):
            values["outcome"] = award.group(0)
    return values


def _looks_like_result(line: str) -> bool:
    return any(marker in line for marker in ("达", "提升", "降低", "完成", "获奖", "SRCC", "准确率"))


def _section_type_for(line: str) -> EvidenceType | None:
    normalized = line.lower().replace(" ", "").rstrip("：:")
    for evidence_type, headings in _SECTION_TYPES:
        if normalized in headings:
            return evidence_type
    return None
